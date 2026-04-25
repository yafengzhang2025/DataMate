# -- encoding: utf-8 --

"""
Description:
Create: 2025/01/18
"""
import random
import string
import time
from io import BytesIO
from typing import List, Union, Dict, Any

import bs4
import docx.table
import numpy as np
import pandas as pd
from docx import Document
from loguru import logger

from datamate.core.base_op import Mapper


class TextToWord(Mapper):
    def __init__(self, *args, **kwargs):
        # 随机生成两个长度超过10的字符串用作标识合并单元格使用
        super().__init__(*args, **kwargs)
        self.delete = ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(15))
        self.abundant = ''.join(random.choice(string.ascii_uppercase + string.digits) for _ in range(15)) + "VVV"

        # 生成字典记录单元格合并信息
        self.diagonal_merge = dict()

    @staticmethod
    def _to_clean_paragraphs(table: docx.table.Table, position: tuple) -> None:
        """删除单元格内多余的换行符"""
        clearn_paragraphs = []
        paragraphs = table.cell(position[0], position[1]).text
        for paragraph in paragraphs:
            clearn_paragraph = paragraph.replace('\n', " ")
            clearn_paragraphs.append(clearn_paragraph)
        table.cell(position[0], position[1]).text = clearn_paragraphs

    @staticmethod
    def _needs_merge(soup: bs4.BeautifulSoup) -> bool:
        """查看是否有合并单元格"""
        all_td = soup.select('tr td')
        for td in all_td:
            if td.has_attr('colspan') or td.has_attr('rowspan'):  # 表示有合并单元格，直接返回True
                return True
        return False

    @staticmethod
    def _find_html_tds(html_tr: bs4.element.Tag) -> bs4.element.ResultSet:
        """判断bs4.element.Tag, 如果以<td>开头需要特殊处理"""
        if str(html_tr).startswith("<td") and str(html_tr).endswith("</td>"):
            html_tds = [html_tr]
        else:
            html_tds = html_tr.find_all('td')
        return html_tds

    def execute(self, sample: Dict[str, Any]) -> Dict[str, Any]:
        """将文本信息转换为docx文件流"""
        start = time.time()
        self.read_file_first(sample)
        sample[self.data_key] = self._txt_to_docx(sample[self.text_key])  # 将文字转换为word字符串流
        sample[self.text_key] = ""
        sample["target_type"] = "docx"
        logger.info(f"fileName: {sample[self.filename_key]}, method: TextToWord costs {time.time() - start:6f} s")
        return sample

    def read_html_with_merge_cell(self, html_table: bs4.BeautifulSoup) -> pd.DataFrame:
        """阅读html文档并根据合并单元格特性进行表内文字去重预处理，以及计算df的行数与列数"""
        html_trs = html_table.find_all('tr')

        # 计算html字段的行数
        row_count = len(html_trs)

        # 计算html字段的列数
        col_count = 0
        for row in html_trs:
            # 处理html字段可能会生成td开头的html文段，需特殊处理
            html_tds = self._find_html_tds(row)

            cur_col = sum([int(html_td['colspan']) if html_td.has_attr('colspan') else 1 for html_td in html_tds])
            col_count = max(col_count, cur_col)

        return self.cell_preprocess(row_count, col_count, html_trs)  # 返回对合并单元格特殊标识后的数据模板

    def mark_span_cell_diagonal(self, parameters: List[int], row_count: int, df: pd.DataFrame) -> pd.DataFrame:
        """将同时具有rowspan和colspan的单元格优先存入字典中，并对需合并单元格进行特殊标注"""
        row, rowspan, col, colspan = parameters

        # 需要动态删减超过行数的rowspan
        min_row = rowspan - 1

        for i in range(row, rowspan):
            if i >= row_count:
                min_row = min(min_row, i - 1)  # 最大rowspan数不能超过模板行数，否则不做后续处理
                break
            for j in range(col, colspan):
                if i == row and j == col:
                    continue
                df.iloc[i, j] = self.delete  # 对需被合并单元格进行特殊标注

        self.diagonal_merge[(row, col)] = (min_row, colspan - 1)  # 将当前单元格以及对应的合并目标单元格存入字典
        return df

    def mark_span_cell_single(self, param: List[int], count: int, df: pd.DataFrame, is_rowspan: bool) -> pd.DataFrame:
        """将具有rowspan或colspan的单元格优先存入字典中，并对需合并单元格进行特殊标注"""
        row, col, span = param

        span_direction = row if is_rowspan else col  # 判断span方向

        min_span = span - 1  # 需要动态删减超过行数的rowspan或colspan

        for i in range(span_direction + 1, span):

            # 如果索引超过行数或列数，更新min_span, 结束循环处理
            if i >= count:
                min_span = min(min_span, i - 1)
                break

            # 根据rowspan或colspan指示判断索引位置
            if is_rowspan:
                df.iloc[i, col] = self.delete  # 对需被合并单元格进行特殊标注
            else:
                df.iloc[row, i] = self.delete

        # 根据rowspan或colspan指示判断索引位置
        if is_rowspan:
            self.diagonal_merge[(row, col)] = (min_span, col)  # 将当前单元格以及对应的合并目标单元格存入字典
        else:
            self.diagonal_merge[(row, col)] = (row, min_span)
        return df

    def mark_abundant_cell_edge(self, param: List[int], df: pd.DataFrame) -> pd.DataFrame:
        """冗余单元格特殊标注"""
        row, col, col_count = param
        if col != col_count:
            for mo in range(col, col_count):
                df.iloc[row, mo] = self.abundant  # 对列中冗余单元格进行特殊标识
        return df

    def cell_preprocess(self, row_count: int, col_count: int, html_trs: List[bs4.element.Tag]) -> pd.DataFrame:
        """判断单元格是否需要被合并，如果需要被合并，call self._mark_abundant_cell 特别标注此单元格"""
        df = pd.DataFrame(np.zeros([row_count, col_count]), dtype=int)

        # 根据网页中的表格，还原在dataframe中，有合并单元格现象的值填在第一个单元格中，其他的用特殊标识填充
        for row in range(row_count):

            # beautifulSoup 处理html表格时可能会生成td开头的html文段，需要特殊处理
            html_tds = self._find_html_tds(html_trs[row])

            # span记录td的索引，td的总数不一定等于col_count, 因为td可能包含colspan
            span = 0

            for col in range(col_count):
                if span >= len(html_tds):
                    df = self.mark_abundant_cell_edge([row, col, col_count], df)  # 标注潜在冗余单元格
                    break

                # 如果框架为空或者字段为特殊标识，则不做后续处理
                if pd.isnull(df.iloc[row, col]) or df.iloc[row, col] == (self.delete or self.abundant):
                    continue
                html_td = html_tds[span]  # 获取单个含有td信息的html tag形式字符串

                # 根据html_td的属性，选择横竖向/横向/竖向的合并预处理
                df = self._choose_span_method(html_td, [row, col, row_count, col_count], df)

                # dataframe当前位置根据td文段赋值
                df.iloc[row, col] = html_td.get_text(strip=True)
                span += 1
        return df

    def _choose_span_method(self, html_td: bs4.element.Tag, param: List[int], df: pd.DataFrame) -> pd.DataFrame:
        """根据信息td属性信息判断合并单元格方向"""
        row, col, row_count, col_count = param

        has_rowspan = html_td.has_attr('rowspan')
        has_colspan = html_td.has_attr('colspan')
        rowspan = int(html_td['rowspan']) if has_rowspan else 0
        colspan = int(html_td['colspan']) if has_colspan else 0

        # 横向与纵向合并的单元格
        if has_colspan and has_rowspan:
            df = self.mark_span_cell_diagonal([row, row + rowspan, col, col + colspan], row_count, df)  # 标注被合并单元格

        # 横向合并的单元格
        elif has_colspan:
            df = self.mark_span_cell_single([row, col, col + colspan], col_count, df, False)

        # 竖向合并的单元格
        elif has_rowspan:
            df = self.mark_span_cell_single([row, col, row + rowspan], row_count, df, True)

        return df

    def _eliminate_values(self, table: docx.table.Table, position: List[int]) -> None:
        """合并单元格前预处理，清除被合并单元格内的特殊标识"""
        try:
            merge_destiny_x, merge_destiny_y = self.diagonal_merge[position]  # 获取此位置指向的目标合并单元格坐标
        except KeyError as e:
            logger.exception(f"Current dictionary is NOT supported: {e}")

        for i in range(position[0], merge_destiny_x + 1):
            for j in range(position[1], merge_destiny_y + 1):
                if i == position[0] and j == position[1]:  # 初始目标不做改动
                    continue
                table.cell(i, j).text = ""  # 清除特殊标识

    def _merge_cell(self, table: docx.table.Table) -> None:
        """合并单元格"""
        for position in self.diagonal_merge.keys():
            merge_destiny_x, merge_destiny_y = self.diagonal_merge[position]  # 获取初始单元格位置指向的目标合并单元格坐标
            self._eliminate_values(table, position)  # 删除特殊标识

            # 如果合并形状不是矩形，则与前单元格合并有冲突
            try:
                table.cell(position[0], position[1]).merge(table.cell(merge_destiny_x, merge_destiny_y))
                self._to_clean_paragraphs(table, position)  # 去除多余的换行符
            except docx.exceptions.InvalidSpanError as e:
                logger.exception(f"Current table cell format is NOT supported: {e}")

    def _get_doc_table(self, dataframe: Union[pd.DataFrame, List], doc: Document, is_merge: bool) -> None:
        """
        dataframe转换为doc表格

        Args:
            dataframe : pd.dataframe 单个文件或文件集
            doc : Python docx 文档
            is_merge : 表示dataframe文件是否存在合并单元格

        """
        if isinstance(dataframe, List):  # dataframe 可能是多个表格
            for data in dataframe:
                self._get_doc_table(data, doc, is_merge)
                return

        rows_num, cols_num = dataframe.shape
        table = doc.add_table(rows=rows_num, cols=cols_num, style="Table Grid")

        for row in range(rows_num):
            cells = table.rows[row].cells
            for col in range(cols_num):

                # 确定单个单元格是否为np.nan
                is_np_nan = pd.isnull(dataframe.iloc[row, col])

                # 确定单个单元格是否为None
                is_none = dataframe.iloc[row, col] is None

                # dataframe.iloc 浮点数固定为0.0
                is_float_zero = dataframe.iloc[row, col] == 0.0

                if not (is_np_nan or is_none or is_float_zero):
                    cells[col].text = str(dataframe.iloc[row, col])
                else:
                    cells[col].text = ""

        # 只有存在合并单元格，才进行合并。否则无需进一步处理
        if is_merge:
            self._merge_cell(table)  # 合并单元格处理

        # 每行列冗余单元格处理
        self._merge_col_abundant_cell(rows_num, cols_num, table)

    def _merge_col_abundant_cell(self, rows_num: int, cols_num: int, table: docx.table.Table) -> None:
        """对word表格每行的冗余单元格进行合并或者清除处理"""
        for row in range(rows_num):
            cells = table.rows[row].cells

            # 双指针记录冗余单元格的起始与完结位置
            start, finish = None, None

            for col in range(cols_num - 1, -1, -1):
                if cells[col].text == self.abundant and not start:  # start 指针赋值，记录第一个冗余单元格位置
                    start = col
                    cells[col].text = ""
                elif cells[col].text == self.abundant and start:  # 冗余单元格消除特殊标识
                    cells[col].text = ""
                elif cells[col].text != self.abundant and start:  # finish指针后赋值，记录冗余单元格序列完结位置
                    finish = col

                # 如果双指针有一端未被赋值，则不进行合并处理
                if not finish or not start:
                    continue

                # 如果合并形状不是矩形，则与前单元格合并有冲突
                try:
                    cells[finish].merge(cells[start])  # 将冗余单元格合入相对最近的单元格
                    self._to_clean_paragraphs(table, (row, finish))  # 去除多余换行符
                except docx.exceptions.InvalidSpanError as e:
                    logger.exception(f"Current table cell format is NOT supported: {e}")

                # 双指针重置
                start, finish = None, None

    def _get_df_with_merge_info(self, html_table: bs4.BeautifulSoup, line: str) -> [pd.DataFrame, bool]:
        """根据是否具有合并单元格属性来确定处理方法"""
        is_merge = True
        if self._needs_merge(html_table):  # 判断html_table是否具有需要合并的单元格
            # 搭建含有合并单元格的表格框架，使用自研算法处理得到数据框架
            df = self.read_html_with_merge_cell(html_table)
        else:
            # 搭建不含有合并单元格的表格框架，使用pd.read_html() 处理得到数据框架
            is_merge = False
            df = pd.read_html(line)
        return df, is_merge

    def _txt_to_docx(self, text: str):
        """将字符串转换为docx文件流"""
        doc = Document()
        for line in text.split("\n"):
            paragraph = doc.add_paragraph()
            try:
                if line.startswith("<html>") and line.endswith("</html>"):  # 一行文字如果以<html>开始并以</html>结束，则判断为是表格
                    self.diagonal_merge = dict()  # 每段html合并单元格信息需重置
                    html_table = bs4.BeautifulSoup(line, 'html.parser')
                    df, is_merge = self._get_df_with_merge_info(html_table, line)  # 得到数据框架
                    self._get_doc_table(df, doc, is_merge)  # 通过数据框架构造word table
                else:  # 一行文字如果是纯文字形式，则是不含表格的文本
                    paragraph.add_run(line)
            except docx.exceptions.InvalidSpanError as e:
                logger.exception(f"Current table cell format is NOT supported: {e}")
            except Exception as e:
                logger.exception(f"Current table shape is Not supported {e}")
        stream = BytesIO()
        doc.save(stream)
        return stream.getvalue()
