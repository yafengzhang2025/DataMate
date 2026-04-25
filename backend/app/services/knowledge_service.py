"""Knowledge base service: create, document upload, indexing, search."""
import asyncio
import json
import os
import uuid
from pathlib import Path

from sqlalchemy import select, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.knowledge import KnowledgeBase, KbDocument


class KnowledgeService:

    @staticmethod
    def _kb_to_dict(kb: KnowledgeBase) -> dict:
        return {
            "id": kb.id,
            "name": kb.name,
            "description": kb.description,
            "embed_model": kb.embed_model,
            "vector_store": kb.vector_store,
            "chunk_strategy": kb.chunk_strategy,
            "chunk_size": kb.chunk_size,
            "chunk_overlap": kb.chunk_overlap,
            "document_count": kb.document_count,
            "chunk_count": kb.chunk_count,
            "created_at": kb.created_at,
        }

    @staticmethod
    def _doc_to_dict(doc: KbDocument) -> dict:
        return {
            "id": doc.id,
            "kb_id": doc.kb_id,
            "name": doc.name,
            "file_type": doc.file_type,
            "file_size": doc.file_size,
            "status": doc.status,
            "chunk_count": doc.chunk_count,
            "error_message": doc.error_message,
            "created_at": doc.created_at,
        }

    async def list_knowledge_bases(self, db: AsyncSession) -> list[dict]:
        rows = (await db.execute(select(KnowledgeBase).order_by(KnowledgeBase.created_at.desc()))).scalars().all()
        return [self._kb_to_dict(r) for r in rows]

    async def get_knowledge_base(self, db: AsyncSession, kb_id: str) -> dict | None:
        kb = await db.get(KnowledgeBase, kb_id)
        return self._kb_to_dict(kb) if kb else None

    async def create_knowledge_base(self, db: AsyncSession, data: dict) -> dict:
        kb_id = str(uuid.uuid4())
        collection_name = f"kb_{kb_id.replace('-', '_')}"
        kb = KnowledgeBase(
            id=kb_id,
            name=data["name"],
            description=data.get("description"),
            embed_model=data.get("embed_model", settings.DEFAULT_EMBED_MODEL),
            vector_store=data.get("vector_store", "chromadb"),
            chunk_strategy=data.get("chunk_strategy", "fixed"),
            chunk_size=data.get("chunk_size", 512),
            chunk_overlap=data.get("chunk_overlap", 64),
            collection_name=collection_name,
        )
        db.add(kb)
        await db.commit()
        await db.refresh(kb)
        return self._kb_to_dict(kb)

    async def delete_knowledge_base(self, db: AsyncSession, kb_id: str) -> bool:
        kb = await db.get(KnowledgeBase, kb_id)
        if not kb:
            return False
        # Delete chromadb collection
        _delete_chroma_collection(kb.collection_name or "")
        await db.delete(kb)
        await db.commit()
        return True

    async def list_documents(self, db: AsyncSession, kb_id: str) -> list[dict]:
        stmt = select(KbDocument).where(KbDocument.kb_id == kb_id).order_by(KbDocument.created_at.desc())
        rows = (await db.execute(stmt)).scalars().all()
        return [self._doc_to_dict(r) for r in rows]

    async def upload_document(
        self, db: AsyncSession, kb_id: str, file_path: str, filename: str, file_size: int
    ) -> dict:
        kb = await db.get(KnowledgeBase, kb_id)
        if not kb:
            raise ValueError("知识库不存在")

        doc_id = str(uuid.uuid4())
        suffix = Path(filename).suffix.lstrip(".").lower()
        dest_dir = Path(settings.DATA_DIR) / "knowledge" / kb_id
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest_path = dest_dir / f"{doc_id}.{suffix}"

        import shutil
        shutil.copy2(file_path, dest_path)

        doc = KbDocument(
            id=doc_id,
            kb_id=kb_id,
            name=filename,
            file_type=suffix,
            file_size=file_size,
            storage_path=str(dest_path),
            status="pending",
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        # Kick off async indexing
        asyncio.create_task(self._index_document(kb_id, doc_id, str(dest_path), suffix))

        return self._doc_to_dict(doc)

    async def delete_document(self, db: AsyncSession, kb_id: str, doc_id: str) -> bool:
        doc = await db.get(KbDocument, doc_id)
        if not doc or doc.kb_id != kb_id:
            return False
        # Remove from vector store
        kb = await db.get(KnowledgeBase, kb_id)
        if kb:
            _delete_doc_from_chroma(kb.collection_name or "", doc_id)
        if doc.storage_path and os.path.exists(doc.storage_path):
            os.remove(doc.storage_path)
        kb_obj = await db.get(KnowledgeBase, kb_id)
        if kb_obj:
            kb_obj.document_count = max(0, kb_obj.document_count - 1)
            kb_obj.chunk_count = max(0, kb_obj.chunk_count - doc.chunk_count)
        await db.delete(doc)
        await db.commit()
        return True

    async def reindex(self, db: AsyncSession, kb_id: str) -> bool:
        kb = await db.get(KnowledgeBase, kb_id)
        if not kb:
            return False
        docs_stmt = select(KbDocument).where(KbDocument.kb_id == kb_id)
        docs = (await db.execute(docs_stmt)).scalars().all()
        for doc in docs:
            doc.status = "pending"
        await db.commit()
        for doc in docs:
            asyncio.create_task(
                self._index_document(kb_id, doc.id, doc.storage_path or "", doc.file_type or "txt")
            )
        return True

    async def search(self, db: AsyncSession, kb_id: str, query: str, top_k: int, threshold: float) -> list[dict]:
        kb = await db.get(KnowledgeBase, kb_id)
        if not kb:
            raise ValueError("知识库不存在")
        return _chroma_search(kb.collection_name or "", query, top_k, threshold)

    async def _index_document(self, kb_id: str, doc_id: str, file_path: str, file_type: str):
        """Background: extract text → chunk → embed → store in chromadb."""
        from app.database import AsyncSessionLocal
        async with AsyncSessionLocal() as db:
            doc = await db.get(KbDocument, doc_id)
            kb = await db.get(KnowledgeBase, kb_id)
            if not doc or not kb:
                return
            try:
                doc.status = "indexing"
                await db.commit()

                text = _extract_text(file_path, file_type)
                chunks = _split_text(text, kb.chunk_size, kb.chunk_overlap)
                _store_chunks(kb.collection_name or "", doc_id, doc.name, chunks)

                doc.status = "indexed"
                doc.chunk_count = len(chunks)
                kb.chunk_count += len(chunks)
                kb.document_count = await _count_docs(db, kb_id)
                await db.commit()
            except Exception as e:
                doc.status = "error"
                doc.error_message = str(e)
                await db.commit()


# ── ChromaDB helpers ─────────────────────────────────────────────────────────

def _get_chroma_client():
    try:
        import chromadb
        return chromadb.PersistentClient(path=settings.CHROMA_PERSIST_DIR)
    except ImportError:
        return None


def _delete_chroma_collection(collection_name: str):
    client = _get_chroma_client()
    if client and collection_name:
        try:
            client.delete_collection(collection_name)
        except Exception:
            pass


def _delete_doc_from_chroma(collection_name: str, doc_id: str):
    client = _get_chroma_client()
    if not client or not collection_name:
        return
    try:
        col = client.get_collection(collection_name)
        col.delete(where={"doc_id": doc_id})
    except Exception:
        pass


def _store_chunks(collection_name: str, doc_id: str, doc_name: str, chunks: list[str]):
    client = _get_chroma_client()
    if not client:
        return
    col = client.get_or_create_collection(collection_name)
    ids = [f"{doc_id}_{i}" for i in range(len(chunks))]
    metadatas = [{"doc_id": doc_id, "document_name": doc_name} for _ in chunks]
    col.add(documents=chunks, ids=ids, metadatas=metadatas)


def _chroma_search(collection_name: str, query: str, top_k: int, threshold: float) -> list[dict]:
    client = _get_chroma_client()
    if not client:
        return []
    try:
        col = client.get_collection(collection_name)
        results = col.query(query_texts=[query], n_results=top_k)
        output = []
        for i, doc in enumerate(results.get("documents", [[]])[0]):
            dist = results.get("distances", [[]])[0][i] if results.get("distances") else 1.0
            score = max(0.0, 1 - dist)  # convert distance to similarity
            if score >= threshold:
                meta = results.get("metadatas", [[]])[0][i] or {}
                output.append({
                    "chunk_id": results["ids"][0][i],
                    "document_name": meta.get("document_name", ""),
                    "content": doc,
                    "score": round(score, 4),
                })
        return output
    except Exception:
        return []


# ── Text extraction ──────────────────────────────────────────────────────────

def _extract_text(file_path: str, file_type: str) -> str:
    if file_type in ("txt", "md"):
        return Path(file_path).read_text(encoding="utf-8")
    if file_type == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        except Exception:
            return ""
    if file_type == "docx":
        try:
            from docx import Document
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""
    if file_type in ("html", "htm"):
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(Path(file_path).read_bytes(), "html.parser")
            return soup.get_text(separator="\n")
        except Exception:
            return ""
    return ""


def _split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    """Simple fixed-size character chunking."""
    if not text.strip():
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return [c for c in chunks if c.strip()]


async def _count_docs(db: AsyncSession, kb_id: str) -> int:
    result = await db.execute(
        select(func.count()).select_from(KbDocument).where(KbDocument.kb_id == kb_id)
    )
    return result.scalar_one()
